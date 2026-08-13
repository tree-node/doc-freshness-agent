"""取り込み（DocumentSource / LocalFolderSource）のテスト。ネットワークに出ない。"""

from __future__ import annotations

from pathlib import Path

import pytest

from app.sources import LocalFolderSource, diff_listings
from app.sources.local import UnsupportedFormatError, detect_format


@pytest.fixture
def watch_root(tmp_path: Path) -> Path:
    (tmp_path / "規定").mkdir()
    (tmp_path / "規定" / "就業規則.md").write_text("# 就業規則\n本文", encoding="utf-8")
    (tmp_path / "メモ_final_v2.txt").write_text("第1条 これは規定です", encoding="utf-8")
    # 対象外が混ざっていても取り込みを止めない
    (tmp_path / "資料.pdf").write_bytes(b"%PDF-1.4")
    (tmp_path / "~$就業規則.md").write_text("Word の一時ファイル", encoding="utf-8")
    return tmp_path


def test_lists_supported_formats_recursively(watch_root: Path) -> None:
    refs = LocalFolderSource(watch_root).list()
    assert [r.doc_id for r in refs] == ["メモ_final_v2.txt", "規定/就業規則.md"]


def test_skips_pdf_and_word_temp_files(watch_root: Path) -> None:
    names = {r.name for r in LocalFolderSource(watch_root).list()}
    assert "資料.pdf" not in names
    assert "~$就業規則.md" not in names


def test_read_returns_text_and_location(watch_root: Path) -> None:
    source = LocalFolderSource(watch_root)
    ref = next(r for r in source.list() if r.format == "md")
    doc = source.read(ref)
    assert doc.text == "# 就業規則\n本文"
    assert doc.doc_id == "規定/就業規則.md"
    assert doc.ref.location.endswith("就業規則.md")


def test_missing_folder_is_reported(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError, match="監視対象フォルダ"):
        LocalFolderSource(tmp_path / "ない").list()


def test_detect_format_rejects_pdf(tmp_path: Path) -> None:
    with pytest.raises(UnsupportedFormatError, match="md/txt/docx"):
        detect_format(tmp_path / "a.pdf")


def test_detect_format_is_case_insensitive(tmp_path: Path) -> None:
    assert detect_format(tmp_path / "A.DOCX") == "docx"


def test_diff_listings_detects_new_updated_removed(watch_root: Path) -> None:
    source = LocalFolderSource(watch_root)
    before = source.list()

    (watch_root / "新規追加.md").write_text("# 新規", encoding="utf-8")
    (watch_root / "メモ_final_v2.txt").write_text("第1条 これは規定です（改訂）", encoding="utf-8")
    (watch_root / "規定" / "就業規則.md").unlink()

    diff = diff_listings(before, source.list())
    assert [r.doc_id for r in diff.added] == ["新規追加.md"]
    assert [r.doc_id for r in diff.updated] == ["メモ_final_v2.txt"]
    assert [r.doc_id for r in diff.removed] == ["規定/就業規則.md"]
    assert not diff.is_empty


def test_diff_listings_empty_when_unchanged(watch_root: Path) -> None:
    refs = LocalFolderSource(watch_root).list()
    assert diff_listings(refs, refs).is_empty


def test_docx_headings_and_table_keep_document_order(tmp_path: Path) -> None:
    """表が「直前の見出し」の下に来ること。python-docx の既定の列挙では順序が失われる。"""
    from docx import Document

    document = Document()
    document.add_heading("育児・介護休業規程", level=1)
    document.add_heading("第3章 子の看護休暇", level=2)
    document.add_paragraph("第16条 労働者は、看護休暇を取得できる。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "区分"
    table.cell(0, 1).text = "日数"
    table.cell(1, 0).text = "子1人"
    table.cell(1, 1).text = "5日"
    document.add_paragraph("以上")
    path = tmp_path / "規程.docx"
    document.save(str(path))

    source = LocalFolderSource(tmp_path)
    ref = next(r for r in source.list() if r.format == "docx")
    lines = source.read(ref).text.splitlines()

    assert lines[0] == "# 育児・介護休業規程"
    assert lines[1] == "## 第3章 子の看護休暇"
    assert lines[2].startswith("第16条")
    assert lines[3] == "区分 | 日数"
    assert lines[4] == "子1人 | 5日"
    assert lines[5] == "以上"
