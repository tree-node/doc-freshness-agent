"""修正版ファイル生成のテスト。ネットワークに出ない。"""

from __future__ import annotations

from pathlib import Path

import pytest

from app.proposal import (
    FixEdit,
    apply_edits_to_text,
    find_span,
    generate_revised_document,
)
from app.sources.local import docx_to_text

BEFORE = "小学校就学の始期に達するまでの子を養育する従業員は、看護休暇を取得できる。"
AFTER = "小学校第三学年修了前の子を養育する従業員は、子の看護等休暇を取得できる。"


def an_edit(before=BEFORE, after=AFTER, location="第21条") -> FixEdit:
    return FixEdit(before=before, after=after, location=location)


# --- 位置の特定 ---------------------------------------------------------------


def test_find_span_matches_exactly() -> None:
    text = f"前文。{BEFORE}後文。"
    start, end = find_span(text, BEFORE)
    assert text[start:end] == BEFORE


def test_find_span_ignores_differences_in_whitespace() -> None:
    """LLMが出す修正前の文は、改行やスペースの入り方が本文と一致しないことがある。"""
    text = "第21条（子の看護休暇）\n小学校就学の始期に達するまでの子を養育する\n従業員は、看護休暇を取得できる。\n"
    span = find_span(text, BEFORE)
    assert span is not None
    assert "".join(text[span[0] : span[1]].split()) == "".join(BEFORE.split())


def test_find_span_returns_none_when_absent() -> None:
    assert find_span("まったく別の本文です。", BEFORE) is None


def test_find_span_ignores_empty_needle() -> None:
    assert find_span("本文", "   ") is None


# --- テキストへの適用 ---------------------------------------------------------


def test_applies_the_replacement() -> None:
    text = f"第21条\n{BEFORE}\n以上。"
    revised, applied, skipped = apply_edits_to_text(text, [an_edit()])
    assert AFTER in revised
    assert BEFORE not in revised
    assert len(applied) == 1
    assert not skipped


def test_reports_instead_of_silently_skipping() -> None:
    """当たらなかった直しは黙って飛ばさない。人間が気づける形で返す。"""
    revised, applied, skipped = apply_edits_to_text("無関係な本文。", [an_edit()])
    assert revised == "無関係な本文。"
    assert not applied
    assert len(skipped) == 1
    assert "見つかりません" in skipped[0][1]


def test_applies_several_edits_without_shifting_each_other() -> None:
    text = "第1条 アアア。\n第2条 イイイ。\n第3条 ウウウ。"
    edits = [
        FixEdit(before="アアア", after="アアアアアアアアア", location="第1条"),
        FixEdit(before="ウウウ", after="ウ", location="第3条"),
    ]
    revised, applied, skipped = apply_edits_to_text(text, edits)
    assert "第1条 アアアアアアアアア。" in revised
    assert "第3条 ウ。" in revised
    assert "第2条 イイイ。" in revised
    assert len(applied) == 2
    assert not skipped


def test_overlapping_edits_are_reported() -> None:
    text = f"{BEFORE}"
    edits = [an_edit(), FixEdit(before="小学校就学の始期", after="小学校第三学年修了前", location="第21条")]
    _, applied, skipped = apply_edits_to_text(text, edits)
    assert len(applied) == 1
    assert len(skipped) == 1
    assert "重なっています" in skipped[0][1]


def test_empty_proposal_is_skipped() -> None:
    _, applied, skipped = apply_edits_to_text("本文", [FixEdit(before="", after="", location="第1条")])
    assert not applied
    assert "空です" in skipped[0][1]


# --- ファイル生成 -------------------------------------------------------------


def test_generates_a_revised_markdown_without_touching_the_original(tmp_path: Path) -> None:
    """**元のファイルは書き換えない**（DESIGN.md 設計原則2）。"""
    source = tmp_path / "就業規則.md"
    original = f"# 就業規則\n\n第21条（子の看護休暇）\n{BEFORE}\n"
    source.write_text(original, encoding="utf-8")

    generated = generate_revised_document(source, "就業規則.md", [an_edit()], tmp_path / "out")

    assert source.read_text(encoding="utf-8") == original  # 元は無傷
    assert AFTER in generated.output_path.read_text(encoding="utf-8")
    assert generated.fully_applied
    assert generated.replace_target == str(source)


def test_keeps_the_folder_structure_of_the_original(tmp_path: Path) -> None:
    source = tmp_path / "規程.md"
    source.write_text(BEFORE, encoding="utf-8")
    generated = generate_revised_document(
        source, "株式会社サクラベース/就業規則.md", [an_edit()], tmp_path / "out"
    )
    assert generated.output_path == tmp_path / "out" / "株式会社サクラベース" / "就業規則.md"


def test_missing_source_is_reported(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError, match="元のファイル"):
        generate_revised_document(tmp_path / "ない.md", "ない.md", [an_edit()], tmp_path / "out")


def build_docx(path: Path, paragraphs: list[str]) -> None:
    from docx import Document

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(str(path))


def test_generates_a_revised_docx(tmp_path: Path) -> None:
    source = tmp_path / "休暇規程.docx"
    build_docx(source, ["第5章 休暇等", f"第21条（子の看護休暇）\n{BEFORE}", "以上"])

    generated = generate_revised_document(source, "休暇規程.docx", [an_edit()], tmp_path / "out")

    revised = docx_to_text(generated.output_path)
    assert AFTER in revised
    assert BEFORE not in revised
    assert "以上" in revised  # 関係ない段落は残る
    assert generated.fully_applied


def test_docx_replacement_keeps_the_rest_of_the_paragraph(tmp_path: Path) -> None:
    """段落を作り直さず、該当箇所だけを差し替える。"""
    source = tmp_path / "規程.docx"
    build_docx(source, [f"前置き。{BEFORE}後置き。"])

    generated = generate_revised_document(source, "規程.docx", [an_edit()], tmp_path / "out")

    revised = docx_to_text(generated.output_path)
    assert revised.startswith("前置き。")
    assert revised.endswith("後置き。")
    assert AFTER in revised


def test_docx_replacement_works_across_runs(tmp_path: Path) -> None:
    """文が書式の切れ目でランに分かれていても当てられること。"""
    from docx import Document

    source = tmp_path / "規程.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("小学校就学の始期に")
    paragraph.add_run("達するまでの子を養育する従業員は、")
    paragraph.add_run("看護休暇を取得できる。")
    document.save(str(source))

    generated = generate_revised_document(source, "規程.docx", [an_edit()], tmp_path / "out")

    assert generated.fully_applied
    assert AFTER in docx_to_text(generated.output_path)


def test_docx_reports_an_edit_that_did_not_match(tmp_path: Path) -> None:
    source = tmp_path / "規程.docx"
    build_docx(source, ["まったく別の本文。"])
    generated = generate_revised_document(source, "規程.docx", [an_edit()], tmp_path / "out")
    assert not generated.applied
    assert len(generated.skipped) == 1
    assert not generated.fully_applied


def test_pdf_is_rejected(tmp_path: Path) -> None:
    """対応形式は md/txt/docx のみ。"""
    from app.sources.local import UnsupportedFormatError

    source = tmp_path / "資料.pdf"
    source.write_bytes(b"%PDF-1.4")
    with pytest.raises(UnsupportedFormatError):
        generate_revised_document(source, "資料.pdf", [an_edit()], tmp_path / "out")
