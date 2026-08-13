#!/usr/bin/env python3
"""デモデータの docx を整形して監視対象フォルダに配置する（再実行可）。

やること:
  - 厚労省モデル就業規則の**解説文**を本文から取り除く。実在の社内規程には無い文で、
    残すと判定に流れてノイズになる（実際に「影響なし」判定が1件それで出ていた）
  - 会社名を架空名で統一する（元データは1ファイルにしか会社名が入っていなかった）
  - 受け取った原本は `demo-data/_原本/` に退避する。**監視対象フォルダには置かない**

    実行: PYTHONPATH=. python scripts/build_demo_data.py

出典: 厚生労働省「モデル就業規則」。政府標準利用規約に基づき利用。
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DEMO = ROOT / "demo-data"
ORIGINALS = DEMO / "_原本"
WATCH = DEMO / "監視対象"

VIOLATION_COMPANY = "株式会社サクラベース"
CLEAN_COMPANY = "サクラベース・ロジスティクス株式会社"
OLD_COMPANY = "株式会社サンプル テック"

# 厚労省モデル就業規則の解説文（実在の規程には無い）
COMMENTARY = re.compile(r"^【第\d+条|ことが必要です。$|^本規程例|^モデル就業規則")

# 原本ファイル名 → (配置先の会社, 文書名)
LAYOUT = [
    ("就業規定＿違反ver.docx", VIOLATION_COMPANY, "就業規則"),
    ("休暇規定_違反ver.docx", VIOLATION_COMPANY, "休暇規程"),
    ("賃金規定_違反ver.docx", VIOLATION_COMPANY, "賃金規程"),
    ("就業規定.docx", CLEAN_COMPANY, "就業規則"),
    ("休暇規定.docx", CLEAN_COMPANY, "休暇規程"),
    ("賃金規定.docx", CLEAN_COMPANY, "賃金規程"),
]


def stash_originals() -> None:
    """受け取ったままの docx を _原本/ へ退避する（監視対象から外すため）。"""
    ORIGINALS.mkdir(parents=True, exist_ok=True)
    for path in sorted(DEMO.glob("*.docx")):
        shutil.move(str(path), str(ORIGINALS / path.name))
        print(f"  退避: {path.name}")
    legacy = DEMO / "違反一覧.txt"
    if legacy.exists():
        shutil.move(str(legacy), str(ORIGINALS / legacy.name))
        print(f"  退避: {legacy.name}（正解表は 答え合わせ/ に作り直す）")


def clean_document(src: Path, dest: Path, company: str, title: str) -> None:
    document = Document(str(src))

    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text and COMMENTARY.search(text):
            paragraph._element.getparent().remove(paragraph._element)

    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, OLD_COMPANY, company)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, OLD_COMPANY, company)

    first = document.paragraphs[0]
    if not first.text.strip().startswith(company):
        first.insert_paragraph_before(f"{company}　{title}")

    dest.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(dest))
    print(f"  {dest.relative_to(DEMO)}")


def _replace_in_paragraph(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    # 文字列が複数のランに分かれている場合があるため、見つからなければ先頭ランに集約する
    if any(old in run.text for run in paragraph.runs):
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
        return
    if paragraph.runs:
        paragraph.runs[0].text = paragraph.text.replace(old, new)
        for run in paragraph.runs[1:]:
            run.text = ""


def main() -> int:
    print("デモデータの docx を整形します")
    stash_originals()

    print("\n規程を配置:")
    for filename, company, title in LAYOUT:
        src = ORIGINALS / filename
        if not src.exists():
            print(f"  スキップ（原本なし）: {filename}")
            continue
        clean_document(src, WATCH / company / f"{title}.docx", company, title)

    # 同一内容の重複配置。判定は1回、起票は両方のファイルに出ることの確認用
    duplicate = WATCH / "共有フォルダ（未整理）" / "就業規則.docx"
    source = WATCH / VIOLATION_COMPANY / "就業規則.docx"
    if source.exists():
        duplicate.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(source, duplicate)
        print(f"  {duplicate.relative_to(DEMO)}（重複配置）")

    print("\n完了。監視対象は demo-data/監視対象/ です。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
