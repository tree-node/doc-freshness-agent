"""修正版ファイルの生成（経路Aの出口）。

**元のファイルは絶対に書き換えない**（DESIGN.md 設計原則2）。別の場所に修正版を作り、
置換先のパスを添えて人間に渡す。適用するかどうかは人間が決める。

置換は機械的に行う。文面を作り直すのはLLMの仕事ではなく、
LLM が出した「この一文をこう直す」（fix_proposal）を、そのとおりに当てるだけ。

うまく当たらなかった箇所は**黙って飛ばさない**。適用できた分と、できなかった分を
理由つきで返す（見逃し側に倒す＝人間が気づける形にする）。
"""

from __future__ import annotations

import shutil
from dataclasses import dataclass, field
from pathlib import Path

from app.sources.local import detect_format, docx_to_text


@dataclass(frozen=True)
class FixEdit:
    """1箇所の直し。LLMが出した before / after をそのまま使う。"""

    before: str
    after: str
    location: str = ""  # 表示用（第28条第2項 など）

    @property
    def is_applicable(self) -> bool:
        return bool(self.before.strip()) and bool(self.after.strip())


@dataclass
class GeneratedDocument:
    """生成した修正版ファイル1つ。"""

    doc_id: str
    output_path: Path
    replace_target: str  # 置換先（元ファイルの場所）。人間に「どこに戻すか」を示す
    applied: list[FixEdit] = field(default_factory=list)
    skipped: list[tuple[FixEdit, str]] = field(default_factory=list)

    @property
    def fully_applied(self) -> bool:
        return not self.skipped

    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "output_path": str(self.output_path),
            "replace_target": self.replace_target,
            "applied": [{"location": e.location, "before": e.before, "after": e.after} for e in self.applied],
            "skipped": [
                {"location": e.location, "before": e.before, "reason": reason} for e, reason in self.skipped
            ],
        }


# --- 空白の違いを吸収した検索 -------------------------------------------------


def _normalize(text: str) -> tuple[str, list[int]]:
    """空白を落とした文字列と、その各文字が元テキストの何文字目かの対応表を返す。

    LLMが出す before は、改行やスペースの入り方が元の文書と一致しないことがある。
    そのまま `str.find` すると当たらないので、空白を除いた上で照合し、
    見つかった位置を元テキストの位置に戻す。
    """
    stripped: list[str] = []
    positions: list[int] = []
    for index, char in enumerate(text):
        if char.isspace():
            continue
        stripped.append(char)
        positions.append(index)
    return "".join(stripped), positions


def find_span(haystack: str, needle: str) -> tuple[int, int] | None:
    """needle が haystack のどこにあるかを返す（空白の違いは無視）。無ければ None。"""
    if not needle.strip():
        return None

    exact = haystack.find(needle)
    if exact >= 0:
        return exact, exact + len(needle)

    flat_hay, positions = _normalize(haystack)
    flat_needle, _ = _normalize(needle)
    if not flat_needle:
        return None
    at = flat_hay.find(flat_needle)
    if at < 0:
        return None
    start = positions[at]
    end = positions[at + len(flat_needle) - 1] + 1
    return start, end


def apply_edits_to_text(text: str, edits: list[FixEdit]) -> tuple[str, list[FixEdit], list[tuple[FixEdit, str]]]:
    """テキストに直しを当てる。当たらなかったものは理由つきで返す。

    後ろの箇所から当てていくことで、前の置換で位置がずれるのを避ける。
    """
    located: list[tuple[int, int, FixEdit]] = []
    skipped: list[tuple[FixEdit, str]] = []

    for edit in edits:
        if not edit.is_applicable:
            skipped.append((edit, "修正案が空です"))
            continue
        span = find_span(text, edit.before)
        if span is None:
            skipped.append((edit, "修正前の文が本文に見つかりませんでした"))
            continue
        if any(not (span[1] <= s or span[0] >= e) for s, e, _ in located):
            skipped.append((edit, "ほかの修正と範囲が重なっています"))
            continue
        located.append((span[0], span[1], edit))

    applied: list[FixEdit] = []
    for start, end, edit in sorted(located, key=lambda item: -item[0]):
        text = text[:start] + edit.after + text[end:]
        applied.append(edit)
    applied.reverse()
    return text, applied, skipped


# --- docx -------------------------------------------------------------------


def _apply_edit_to_paragraph(paragraph, edit: FixEdit) -> bool:
    """段落の中の該当箇所だけを差し替える。

    文が複数のラン（書式の切れ目）にまたがっていることがあるので、該当する範囲の
    先頭ランに置換後の文を入れ、残りのランを空にする。段落ごと作り直さないのは、
    同じ段落の関係ない部分の書式を壊さないため。
    """
    span = find_span(paragraph.text, edit.before)
    if span is None:
        return False
    start, end = span

    cursor = 0
    first = True
    for run in paragraph.runs:
        run_start, run_end = cursor, cursor + len(run.text)
        cursor = run_end
        if run_end <= start or run_start >= end:
            continue  # この範囲は置換対象の外
        head = run.text[: max(start - run_start, 0)]
        tail = run.text[max(end - run_start, 0) :] if run_end > end else ""
        run.text = f"{head}{edit.after}{tail}" if first else f"{head}{tail}"
        first = False
    return not first


def _apply_edits_to_docx(source: Path, dest: Path, edits: list[FixEdit]) -> tuple[list[FixEdit], list[tuple[FixEdit, str]]]:
    from docx import Document

    document = Document(str(source))
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    applied: list[FixEdit] = []
    skipped: list[tuple[FixEdit, str]] = []

    for edit in edits:
        if not edit.is_applicable:
            skipped.append((edit, "修正案が空です"))
            continue
        for paragraph in paragraphs:
            if _apply_edit_to_paragraph(paragraph, edit):
                applied.append(edit)
                break
        else:
            skipped.append((edit, "修正前の文が本文に見つかりませんでした"))

    dest.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(dest))
    return applied, skipped


# --- 入口 --------------------------------------------------------------------


def generate_revised_document(
    source_path: Path, doc_id: str, edits: list[FixEdit], output_dir: Path
) -> GeneratedDocument:
    """修正版ファイルを作る。元のファイルには触らない。"""
    if not source_path.exists():
        raise FileNotFoundError(f"元のファイルが見つかりません: {source_path}")

    fmt = detect_format(source_path)
    dest = output_dir / doc_id
    dest.parent.mkdir(parents=True, exist_ok=True)

    if fmt == "docx":
        applied, skipped = _apply_edits_to_docx(source_path, dest, edits)
    else:
        text = source_path.read_text(encoding="utf-8")
        revised, applied, skipped = apply_edits_to_text(text, edits)
        dest.write_text(revised, encoding="utf-8")

    generated = GeneratedDocument(
        doc_id=doc_id, output_path=dest, replace_target=str(source_path), applied=applied, skipped=skipped
    )
    _verify(dest, fmt, generated)
    return generated


def _verify(dest: Path, fmt: str, generated: GeneratedDocument) -> None:
    """当てたはずの文が本当に入っているかを、生成後のファイルから読み直して確かめる。

    docx はランの分かれ方によって置換が空振りすることがある。生成できたつもりで
    中身が変わっていない、という事故を防ぐ。
    """
    if not generated.applied:
        return
    text = docx_to_text(dest) if fmt == "docx" else dest.read_text(encoding="utf-8")
    still_missing = [edit for edit in generated.applied if find_span(text, edit.after) is None]
    for edit in still_missing:
        generated.applied.remove(edit)
        generated.skipped.append((edit, "置換したはずの文が生成後のファイルに見つかりません"))


def edits_from_result(result: dict, doc_id: str) -> list[FixEdit]:
    """検知結果JSONから、その文書に対する直しを取り出す。

    起票（alerts）は「影響あり」だけを載せているので、そこから拾えば
    影響なしと判定された箇所を誤って直すことはない。
    """
    edits: list[FixEdit] = []
    seen: set[tuple[str, str]] = set()
    for alert in result.get("alerts", []):
        if alert.get("doc_id") != doc_id:
            continue
        finding = alert.get("finding") or {}
        proposal = finding.get("fix_proposal") or {}
        before, after = proposal.get("before"), proposal.get("after")
        if not before or not after:
            continue
        key = (before, after)
        if key in seen:
            continue  # 同じ内容のチャンクが複数あると同じ直しが複数回来る
        seen.add(key)
        edits.append(FixEdit(before=before, after=after, location=finding.get("evidence_location") or ""))
    return edits


def affected_doc_ids(result: dict) -> list[str]:
    """その検知結果で「要対応」と判定された文書。"""
    seen: list[str] = []
    for alert in result.get("alerts", []):
        doc_id = alert.get("doc_id")
        if doc_id and doc_id not in seen:
            seen.append(doc_id)
    return seen


def copy_untouched(source_path: Path, doc_id: str, output_dir: Path) -> Path:
    """直すところが無い文書をそのまま複製する（まとめてDLするとき用）。"""
    dest = output_dir / doc_id
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(source_path, dest)
    return dest
