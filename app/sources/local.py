"""ローカルフォルダからの取り込み。

docx はプレーンテキストに変換して扱う。**段落を改行1つで連結する**ことを規約とし、
チャンクの文字オフセットはこの変換後テキスト上の位置とする（修正版ファイル生成時に
段落番号へ戻せるようにするため、変換規則を勝手に変えないこと）。
"""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

from app.sources.base import (
    SUPPORTED_FORMATS,
    DocumentFormat,
    DocumentRef,
    DocumentSource,
    SourceDocument,
)

# 表のセル区切り。md のテーブル記法に寄せて後段のチャンク分割で表として拾えるようにする
DOCX_CELL_SEPARATOR = " | "


class UnsupportedFormatError(ValueError):
    pass


def detect_format(path: Path) -> DocumentFormat:
    suffix = path.suffix.lower().lstrip(".")
    if suffix in SUPPORTED_FORMATS:
        return suffix  # type: ignore[return-value]
    raise UnsupportedFormatError(f"対応形式は md/txt/docx のみです: {path.name}")


def docx_to_text(path: Path) -> str:
    """docx を段落・表のテキストに変換する。

    見出しスタイルの情報はチャンク分割で必要なので、見出し段落は md 見出し記法に変換して
    残す（分割ルールを md と docx で共通化するため）。
    """
    from docx import Document  # 重いので必要になったときだけ読む

    document = Document(str(path))
    lines: list[str] = []

    for block in _iter_docx_blocks(document):
        if block[0] == "paragraph":
            paragraph = block[1]
            text = paragraph.text.strip()
            if not text:
                continue
            level = _heading_level(paragraph)
            lines.append(f"{'#' * level} {text}" if level else text)
        else:
            for row in block[1].rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                lines.append(DOCX_CELL_SEPARATOR.join(cells))

    return "\n".join(lines)


def _iter_docx_blocks(document):  # type: ignore[no-untyped-def]
    """段落と表を**本文の出現順**で列挙する。

    python-docx の document.paragraphs / document.tables は別々の列挙になっていて
    順序が失われる（表が本文のどの見出しの下にあるか分からなくなる）ため、
    XML の並びをたどる。
    """
    from docx.document import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    parent = document if isinstance(document, DocxDocument) else document._parent
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield ("paragraph", Paragraph(child, parent))
        elif child.tag.endswith("}tbl"):
            yield ("table", Table(child, parent))


def _heading_level(paragraph) -> int:  # type: ignore[no-untyped-def]
    """見出しスタイルなら 1〜3 を返す。日本語版Wordの「見出し 1」にも対応する。"""
    style = (paragraph.style.name or "") if paragraph.style else ""
    for token in ("Heading", "見出し"):
        if style.startswith(token):
            tail = style[len(token) :].strip()
            if tail.isdigit():
                return min(int(tail), 3)
    if style in ("Title", "表題"):
        return 1
    return 0


class LocalFolderSource(DocumentSource):
    """フォルダ配下を再帰的に取り込む（登録単位はフォルダ）。"""

    def __init__(self, root: Path) -> None:
        self.root = Path(root)

    @property
    def source_id(self) -> str:
        return str(self.root)

    def list(self) -> list[DocumentRef]:
        if not self.root.exists():
            raise FileNotFoundError(f"監視対象フォルダがありません: {self.root}")

        refs: list[DocumentRef] = []
        for path in sorted(self.root.rglob("*")):
            if not path.is_file() or path.name.startswith("~$"):
                continue
            try:
                fmt = detect_format(path)
            except UnsupportedFormatError:
                # 対象外の形式は静かに飛ばす（PDF等が混在していても取り込みは止めない）
                continue
            stat = path.stat()
            refs.append(
                DocumentRef(
                    doc_id=path.relative_to(self.root).as_posix(),
                    name=path.name,
                    location=str(path),
                    format=fmt,
                    size=stat.st_size,
                    modified_at=datetime.fromtimestamp(stat.st_mtime, UTC).isoformat(
                        timespec="seconds"
                    ),
                )
            )
        return refs

    def read(self, ref: DocumentRef) -> SourceDocument:
        path = self.root / ref.doc_id
        if ref.format == "docx":
            text = docx_to_text(path)
        else:
            text = path.read_text(encoding="utf-8")
        return SourceDocument(ref=ref, text=text)
